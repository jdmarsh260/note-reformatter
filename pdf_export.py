import re
from docx import Document


# creates functions to be used later on in code
def replace_last(string, find, replace):
    reversed = string[::-1]
    replaced = reversed.replace(find[::-1], replace[::-1], 1)
    return replaced[::-1]

def docx_find_replace_text(doc, search_text, replace_text):
    paragraphs = list(doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraphs.append(paragraph)
    for p in paragraphs:
        if search_text in p.text:
            inline = p.runs
            # Replace strings and retain the same style.
            # The text to be replaced can be split over several runs so
            # search through, identify which runs need to have text replaced
            # then replace the text in those identified
            started = False
            search_index = 0
            # found_runs is a list of (inline index, index of match, length of match)
            found_runs = list()
            found_all = False
            replace_done = False
            for i in range(len(inline)):

                # case 1: found in single run so short circuit the replace
                if search_text in inline[i].text and not started:
                    found_runs.append((i, inline[i].text.find(search_text), len(search_text)))
                    text = inline[i].text.replace(search_text, str(replace_text))
                    inline[i].text = text
                    replace_done = True
                    found_all = True
                    break


#ask user for name of text file
txt_file = input("Enter the path for the txt document you wish to reformat:")


# writes a final newline to text as a precaution to make sure regular statements will work as intended
f = open(txt_file, 'a')
f.write('\n')


# saves the contents of the text file as a string
f = open(txt_file, 'r', encoding = "utf8")
contents = f.read()
f.close()


# reformats most of the issues in the text
contents = re.sub(r'\u000C', '', contents)
contents = contents.replace('- ', '')
contents = re.sub(r'\n(\s)+\n', '\n\n', contents)
contents = contents.replace('[BLANK PAGE BLANK PAGE BLANK PAGE]', 'BLANK_PAGE_')
contents = re.sub(r'BLANK_PAGE_ ', 'BLANK_PAGE_', contents)
contents = contents.replace('. \n', '.\n')
contents = re.sub(r'([\.?!]) \n', r'\1\n', contents)
contents = re.sub(r'([^\.?!](\"|\'|“|”)*)\nBLANK_PAGE_\n', r'\1', contents)
contents = contents.replace('BLANK_PAGE_\n', '')
contents = contents.replace('BLANK_PAGE_', '')
contents = contents.replace('\n\n', '\n')
contents = contents.replace('\n', ' [end_of_paragraph].\n')


# splits each sentence into its own line
contents = re.sub(r'([a-z][\.?!](\"|\'|“|”)*)([A-Z])', r'\1 \3', contents)
contents = re.sub(r'([a-z][\.?!](\"|\'|“|”)*)[0-9]+ ((\"|\'|“|”)*[A-Z])', r'\1 \3', contents)
contents = re.sub(r'(.*?[^A-Z\.]{2}[\.?!](\”|\')*)(\s)((\"|\'|“|”)*)([A-Z0-1])', r'\1\n\6', contents)



# sets up PARAGRAPH for each paragraph item
contents = contents.replace(' [end_of_paragraph].\n', '\nPARAGRAPH\n')
contents = replace_last(contents, '\nPARAGRAPH\n','')


# writes the edited contents string into a new text document
f = open('edits.txt', 'w', encoding = "utf8")
f.write(contents)
f.close


# opens that text document and exports each line as an element in a list
f = open('edits.txt', encoding = "utf8")
lines = f.readlines()
contents = []

for line in lines:
    contents.append(line)
f.close()


# fixes troublesome text that may have been incorrectly split
for i in contents:

    if 'Mrs.' in i: # and any other troublesome text
        index_1 = int(contents.index(i))
        index_2 = int(contents.index(i)) + 1

        new_element = contents[index_1]
        new_element = new_element.replace('\n', ' ')
        contents[index_1] = new_element

        new_element = ''.join(contents[x] for x in [index_1, index_2])
        contents[index_1] = new_element
        del contents[index_2]

    else:
        continue


# creates final Word document and adds first line of text to make sure things work as intended
document = Document()
p = document.add_paragraph('', style='List Bullet')
p.add_run('PARAGRAPH').bold = True


# loops through each line saved in the contents list and either formats it as a paragraph or as a sentence under a paragraph
for i in contents:

    if i == 'PARAGRAPH\n':
        p = document.add_paragraph('', style='List Bullet')
        p.add_run(i).bold = True

    elif i != 'PARAGRAPH\n':
        p = document.add_paragraph(i, style='List Bullet 3')


# reformats paragraph line to make it more visually appealing
docx_find_replace_text(document, 'PARAGRAPH\n', 'PARAGRAPH')


# saves final document
final_doc_name = txt_file + '_final.docx'
document.save(final_doc_name)
